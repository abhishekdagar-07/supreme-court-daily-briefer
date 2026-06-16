"""Builds the combined daily Word brief (one .docx for the whole day)."""
import datetime
import logging

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

log = logging.getLogger("worddoc")

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)
SECTIONS = [("Issues", "issues"), ("Facts", "facts"), ("Holding", "holding"),
            ("Ratio", "ratio"), ("Final Order", "final_order")]


def _as_text(value) -> list[str]:
    """Normalise a field that may be a string or a list into paragraphs."""
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value).strip()] if str(value).strip() else ["Not stated in judgment"]


def _heading(doc, text, size, color=NAVY, bold=True, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def _meta_line(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    rl = p.add_run(f"{label}: ")
    rl.bold = True
    rl.font.size = Pt(10)
    rv = p.add_run(str(value))
    rv.font.size = Pt(10)


def build_document(date: datetime.date, items: list[dict]) -> "object":
    """items: list of dicts with keys: meta (row dict), headnote (dict), filename.
    Returns the path to the saved .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Title ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Supreme Court of India")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Daily Judgments Brief")
    rs.font.size = Pt(13)
    rs.font.color.rgb = GREY
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run(date.strftime("%A, %d %B %Y") + f"   •   {len(items)} judgment(s)")
    rd.font.size = Pt(11)
    rd.bold = True

    # ── Contents ──
    _heading(doc, "Contents", 14, space_before=14)
    for i, it in enumerate(items, 1):
        m, h = it["meta"], it["headnote"]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(f"{i}. {m.get('parties','')}")
        rn.bold = True
        rn.font.size = Pt(10.5)
        extra = []
        if m.get("citation"):
            extra.append(m["citation"])
        if m.get("subject") or h.get("subject"):
            extra.append(h.get("subject", ""))
        if extra:
            re_ = p.add_run("  [" + " | ".join(x for x in extra if x) + "]")
            re_.font.size = Pt(9)
            re_.font.color.rgb = GREY
        one = _as_text(h.get("one_line_holding"))[0]
        pl = doc.add_paragraph()
        pl.paragraph_format.left_indent = Inches(0.3)
        pl.paragraph_format.space_after = Pt(6)
        rl = pl.add_run(one)
        rl.italic = True
        rl.font.size = Pt(9.5)
        rl.font.color.rgb = GREY

    # ── Each judgment ──
    for i, it in enumerate(items, 1):
        m, h = it["meta"], it["headnote"]
        doc.add_page_break()
        _heading(doc, f"{i}. {m.get('parties','')}", 14)
        _meta_line(doc, "Neutral Citation", m.get("citation"))
        _meta_line(doc, "Case Number", m.get("case_number"))
        _meta_line(doc, "Diary Number", m.get("diary"))
        _meta_line(doc, "Date of Judgment", m.get("judgment_date"))
        _meta_line(doc, "Bench", m.get("bench"))
        _meta_line(doc, "Judgment authored by", m.get("judgment_by"))
        _meta_line(doc, "Counsel", m.get("advocates"))
        _meta_line(doc, "Subject", h.get("subject"))
        if it.get("filename"):
            _meta_line(doc, "PDF file", it["filename"])
        for title, key in SECTIONS:
            _heading(doc, title, 11.5, space_before=8, space_after=2)
            for para in _as_text(h.get(key)):
                bp = doc.add_paragraph(para)
                bp.paragraph_format.space_after = Pt(3)

    out = config.date_folder(date) / f"Brief {date.strftime('%Y-%m-%d')}.docx"
    doc.save(str(out))
    log.info("saved Word brief: %s", out)
    return out
